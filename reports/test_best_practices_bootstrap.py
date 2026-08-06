import zipfile
from pathlib import Path
from unittest.mock import patch

from django.test import SimpleTestCase, TestCase

from .models import (
    ResourceKnowledgeChunk,
    ResourceKnowledgeDocument,
    ResourceKnowledgeItem,
    ResourceKnowledgeSection,
)
from .resource_knowledge_extraction_service import parse_resource_document
from .resource_knowledge_index_service import index_resource, preview_library
from .resource_library import ResourceFile


TEST_TEMP_ROOT = Path(__file__).resolve().parent / ".test-runtime"


def test_path(case_name: str, file_name: str) -> Path:
    directory = TEST_TEMP_ROOT / case_name
    directory.mkdir(parents=True, exist_ok=True)
    return directory / file_name


def resource_for(path: Path) -> ResourceFile:
    return ResourceFile(
        id="best-practice-test",
        title=path.stem,
        filename=path.name,
        extension=path.suffix.lstrip(".").upper(),
        section="Maintenance & Repair",
        category="Preventive Maintenance",
        level="Tactical",
        folder_path="Best Practices",
        relative_path=f"Best Practices/{path.name}",
        size=path.stat().st_size,
        size_label=f"{path.stat().st_size} B",
        mime_type="text/plain",
        view_url="/resources/test/",
        raw_url="/resources/test/file/",
        is_pdf=path.suffix.lower() == ".pdf",
        is_text=path.suffix.lower() in {".txt", ".md"},
    )


class ResourceParserTests(SimpleTestCase):
    def test_text_parser_preserves_heading_and_content(self):
        path = test_path("text-parser", "practice.md")
        path.write_text(
            "# Preventive Maintenance\n\nInspect the machine before every planned service.",
            encoding="utf-8",
        )
        parsed = parse_resource_document(path)

        self.assertEqual(parsed.parser_name, "text")
        self.assertEqual(parsed.pages[0].number, 1)
        self.assertIn("Preventive Maintenance", parsed.pages[0].text)
        self.assertGreaterEqual(len(parsed.sections), 1)

    def test_docx_parser_extracts_paragraphs_and_tables(self):
        from docx import Document

        path = test_path("docx-parser", "practice.docx")
        document = Document()
        document.add_heading("Maintenance Planning", level=1)
        document.add_paragraph("Review the work scope before shutdown.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Check"
        table.cell(0, 1).text = "Frequency"
        table.cell(1, 0).text = "Oil"
        table.cell(1, 1).text = "Daily"
        document.save(path)
        parsed = parse_resource_document(path)

        self.assertEqual(parsed.parser_name, "python-docx")
        self.assertEqual(len(parsed.tables), 1)
        self.assertIn("Maintenance Planning", parsed.pages[0].text)

    def test_pptx_parser_works_without_external_pptx_dependency(self):
        slide_xml = """<?xml version="1.0" encoding="UTF-8"?>
        <p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
          <p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r><a:t>Review repeat failures</a:t></a:r></a:p>
          </p:txBody></p:sp></p:spTree></p:cSld>
        </p:sld>"""
        path = test_path("pptx-parser", "practice.pptx")
        with zipfile.ZipFile(path, "w") as archive:
            archive.writestr("ppt/slides/slide1.xml", slide_xml)
        parsed = parse_resource_document(path)

        self.assertEqual(parsed.parser_name, "pptx-xml")
        self.assertIn("Review repeat failures", parsed.pages[0].text)

    def test_xlsx_parser_preserves_sheet_as_table(self):
        from openpyxl import Workbook

        path = test_path("xlsx-parser", "practice.xlsx")
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Inspection"
        sheet.append(["Step", "Action"])
        sheet.append([1, "Inspect before service"])
        workbook.save(path)
        parsed = parse_resource_document(path)

        self.assertEqual(parsed.parser_name, "openpyxl")
        self.assertEqual(parsed.tables[0].title, "Inspection")
        self.assertIn("Inspect before service", parsed.tables[0].raw_representation)


class DeterministicBootstrapTests(TestCase):
    def test_preview_has_zero_api_cost_and_writes_nothing(self):
        with patch(
            "reports.resource_knowledge_index_service.list_best_practice_resources",
            return_value=[],
        ):
            result = preview_library()

        self.assertEqual(result["expected_openai_calls"], 0)
        self.assertEqual(result["expected_api_cost"], 0)
        self.assertEqual(ResourceKnowledgeDocument.objects.count(), 0)

    def test_apply_creates_traceable_chunks_without_openai(self):
        path = test_path("apply", "practice.txt")
        path.write_text(
            "PREVENTIVE MAINTENANCE\n\nInspect the equipment before scheduled service. "
            "Ensure all required parts are available.",
            encoding="utf-8",
        )
        resource = resource_for(path)
        with patch(
            "reports.resource_knowledge_index_service.get_resource_path",
            return_value=path,
        ), patch(
            "reports.resource_knowledge_index_service.extract_chunk_knowledge"
        ) as openai_extract:
            result = index_resource(resource)

        self.assertEqual(result["status"], "Indexed")
        self.assertEqual(result["embeddings"], 0)
        openai_extract.assert_not_called()
        document = ResourceKnowledgeDocument.objects.get(resource_id=resource.id)
        chunk = ResourceKnowledgeChunk.objects.get(document=document)
        self.assertEqual(document.metadata_json["openai_calls"], 0)
        self.assertEqual(chunk.extraction_metadata["resource_category"], "Best Practices")
        self.assertTrue(chunk.source_reference)
        self.assertEqual(ResourceKnowledgeSection.objects.filter(document=document).count(), 1)
        item = ResourceKnowledgeItem.objects.get(document=document, is_active=True)
        self.assertEqual(item.extraction_source, "Best Practice Resource")
        self.assertEqual(item.validation_status, "To Review")
        self.assertIn("Inspect the equipment", item.source_excerpt)

    def test_second_identical_apply_is_idempotent(self):
        path = test_path("idempotence", "practice.txt")
        path.write_text("Inspect the machine before service.", encoding="utf-8")
        resource = resource_for(path)
        with patch(
            "reports.resource_knowledge_index_service.get_resource_path",
            return_value=path,
        ):
            index_resource(resource)
            second = index_resource(resource)

        self.assertEqual(second["status"], "skipped")
        self.assertEqual(ResourceKnowledgeDocument.objects.count(), 1)

    def test_reprocess_does_not_deactivate_validated_knowledge(self):
        path = test_path("validated-protection", "practice.txt")
        path.write_text("Inspect the machine before service.", encoding="utf-8")
        resource = resource_for(path)
        with patch(
            "reports.resource_knowledge_index_service.get_resource_path",
            return_value=path,
        ):
            index_resource(resource)
            validated = ResourceKnowledgeItem.objects.get(document__resource_id=resource.id)
            validated.validation_status = "Validated"
            validated.save(update_fields=["validation_status", "updated_at"])
            path.write_text(
                "Inspect the machine before service. Review repeat failures.",
                encoding="utf-8",
            )
            index_resource(resource, force=True)

        validated.refresh_from_db()
        self.assertTrue(validated.is_active)
        self.assertEqual(validated.validation_status, "Validated")
