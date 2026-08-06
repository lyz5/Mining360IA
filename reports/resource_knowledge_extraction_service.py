from __future__ import annotations

import hashlib
import csv
import re
import unicodedata
import zipfile
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from xml.etree import ElementTree


PARSER_VERSION = "2.0"
DEFAULT_MAX_CHUNK_TOKENS = 1500
DEFAULT_MIN_CHUNK_TOKENS = 150


@dataclass(frozen=True)
class ExtractedPage:
    number: int
    text: str


@dataclass(frozen=True)
class ExtractedSection:
    title: str
    content: str
    page_start: int
    page_end: int
    level: int = 1
    section_number: str = ""
    sort_order: int = 0
    heading_path: list[str] = field(default_factory=list)


@dataclass(frozen=True)
class ExtractedTable:
    page: int
    title: str
    headers: list[str]
    rows: list[list[str]]
    raw_representation: str


@dataclass(frozen=True)
class ExtractedVisual:
    page: int
    caption: str = ""
    asset_type: str = "Image"
    extraction_status: str = "Metadata Only"


@dataclass(frozen=True)
class ParsedResourceDocument:
    pages: list[ExtractedPage]
    sections: list[ExtractedSection]
    tables: list[ExtractedTable]
    visuals: list[ExtractedVisual]
    parser_name: str
    parser_version: str = PARSER_VERSION
    language: str = "en"
    ocr_required: bool = False


@dataclass(frozen=True)
class ExtractedChunk:
    index: int
    page_start: int | None
    page_end: int | None
    heading: str
    content: str
    content_hash: str
    token_count: int
    heading_path: list[str] = field(default_factory=list)
    chunk_type: str = "Text"
    source_reference: str = ""


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def normalize_extracted_text(value: str) -> str:
    text = unicodedata.normalize("NFKC", str(value or "")).replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalized_search_text(value: str) -> str:
    return normalize_extracted_text(value).casefold()


def _heading(line: str) -> bool:
    value = line.strip()
    if not value or len(value) > 160:
        return False
    if re.match(r"^(?:\d+(?:\.\d+)*|[A-Z])(?:[.)]|\s+-|\s+)\s*", value):
        return True
    letters = [char for char in value if char.isalpha()]
    return bool(letters) and len(value.split()) <= 14 and (
        value.isupper() or value.istitle()
    )


def _section_number(title: str) -> str:
    match = re.match(r"^\s*(\d+(?:\.\d+)*)", title)
    return match.group(1) if match else ""


def _logical_sections(page: ExtractedPage) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    heading = ""
    content: list[str] = []
    for raw_line in page.text.splitlines():
        line = raw_line.strip()
        if not line:
            if content and content[-1] != "":
                content.append("")
            continue
        if _heading(line) and content:
            body = normalize_extracted_text("\n".join(content))
            if body:
                sections.append((heading, body))
            heading = line
            content = []
        elif _heading(line) and not content:
            heading = line
        else:
            content.append(line)
    body = normalize_extracted_text("\n".join(content))
    if body:
        sections.append((heading, body))
    return sections or [("", page.text)]


def sections_from_pages(pages: list[ExtractedPage]) -> list[ExtractedSection]:
    result: list[ExtractedSection] = []
    order = 0
    for page in pages:
        for heading, content in _logical_sections(page):
            result.append(ExtractedSection(
                title=heading,
                content=content,
                page_start=page.number,
                page_end=page.number,
                level=max(1, (_section_number(heading).count(".") + 1) if heading else 1),
                section_number=_section_number(heading),
                sort_order=order,
                heading_path=[heading] if heading else [],
            ))
            order += 1
    return result


class ResourceDocumentParser:
    name = "base"
    extensions: set[str] = set()

    def supports(self, path: Path) -> bool:
        return path.suffix.casefold() in self.extensions

    def parse(self, path: Path) -> ParsedResourceDocument:
        raise NotImplementedError


class PDFResourceParser(ResourceDocumentParser):
    name = "pypdf"
    extensions = {".pdf"}

    def parse(self, path: Path) -> ParsedResourceDocument:
        indexed_pages = self._indexed_pages(path)
        if indexed_pages is not None:
            return ParsedResourceDocument(
                pages=indexed_pages,
                sections=sections_from_pages(indexed_pages),
                tables=[],
                visuals=[],
                parser_name="local-pdf-text-index",
                ocr_required=not indexed_pages,
            )
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages: list[ExtractedPage] = []
        visuals: list[ExtractedVisual] = []
        for index, page in enumerate(reader.pages, start=1):
            text = normalize_extracted_text(page.extract_text() or "")
            if text:
                pages.append(ExtractedPage(index, text))
            try:
                image_count = len(page.images)
            except Exception:
                image_count = 0
            visuals.extend(ExtractedVisual(index) for _ in range(image_count))
        return ParsedResourceDocument(
            pages=pages,
            sections=sections_from_pages(pages),
            tables=[],
            visuals=visuals,
            parser_name=self.name,
            ocr_required=bool(reader.pages and not pages),
        )

    @staticmethod
    @lru_cache(maxsize=4)
    def _manifest(index_root: Path) -> dict[str, str]:
        manifest_path = index_root / "manifest.csv"
        if not manifest_path.exists():
            return {}
        with manifest_path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
            return {
                str(row.get("relative_path") or "").replace("\\", "/").casefold(): str(row.get("text_file") or "")
                for row in csv.DictReader(handle)
                if str(row.get("status") or "").casefold() == "ok"
            }

    def _indexed_pages(self, path: Path) -> list[ExtractedPage] | None:
        resource_root = next(
            (parent for parent in path.parents if (parent / "_pdf_text_index" / "manifest.csv").exists()),
            None,
        )
        if not resource_root:
            return None
        relative = path.relative_to(resource_root).as_posix().casefold()
        index_root = resource_root / "_pdf_text_index"
        text_file = self._manifest(index_root).get(relative)
        if not text_file:
            return None
        text_path = index_root / text_file
        if not text_path.exists():
            return None
        raw = text_path.read_text(encoding="utf-8", errors="replace")
        matches = list(re.finditer(r"(?m)^===== PAGE (\d+) =====\s*$", raw))
        pages: list[ExtractedPage] = []
        for index, match in enumerate(matches):
            end = matches[index + 1].start() if index + 1 < len(matches) else len(raw)
            text = normalize_extracted_text(raw[match.end():end])
            if text:
                pages.append(ExtractedPage(int(match.group(1)), text))
        return pages


class DOCXResourceParser(ResourceDocumentParser):
    name = "python-docx"
    extensions = {".docx"}

    def parse(self, path: Path) -> ParsedResourceDocument:
        from docx import Document

        document = Document(str(path))
        lines: list[str] = []
        sections: list[ExtractedSection] = []
        current_title = ""
        current_content: list[str] = []
        order = 0

        def flush():
            nonlocal current_content, order
            content = normalize_extracted_text("\n".join(current_content))
            if content:
                sections.append(ExtractedSection(
                    title=current_title,
                    content=content,
                    page_start=1,
                    page_end=1,
                    level=1,
                    section_number=_section_number(current_title),
                    sort_order=order,
                    heading_path=[current_title] if current_title else [],
                ))
                order += 1
            current_content = []

        for paragraph in document.paragraphs:
            text = normalize_extracted_text(paragraph.text)
            if not text:
                continue
            lines.append(text)
            style_name = str(getattr(paragraph.style, "name", "") or "")
            if style_name.casefold().startswith("heading"):
                flush()
                current_title = text
            else:
                current_content.append(text)
        flush()
        tables: list[ExtractedTable] = []
        for index, table in enumerate(document.tables, start=1):
            rows = [[normalize_extracted_text(cell.text) for cell in row.cells] for row in table.rows]
            rows = [row for row in rows if any(row)]
            if not rows:
                continue
            raw = "\n".join(" | ".join(row) for row in rows)
            tables.append(ExtractedTable(1, f"Table {index}", rows[0], rows[1:], raw))
            lines.append(raw)
        pages = [ExtractedPage(1, normalize_extracted_text("\n\n".join(lines)))] if lines else []
        return ParsedResourceDocument(
            pages=pages,
            sections=sections or sections_from_pages(pages),
            tables=tables,
            visuals=[],
            parser_name=self.name,
            ocr_required=not pages,
        )


class PPTXResourceParser(ResourceDocumentParser):
    name = "pptx-xml"
    extensions = {".pptx"}
    namespaces = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}

    def parse(self, path: Path) -> ParsedResourceDocument:
        pages: list[ExtractedPage] = []
        with zipfile.ZipFile(path) as archive:
            slide_names = sorted(
                (
                    name for name in archive.namelist()
                    if re.match(r"ppt/slides/slide\d+\.xml$", name)
                ),
                key=lambda value: int(re.search(r"(\d+)", Path(value).stem).group(1)),
            )
            for number, name in enumerate(slide_names, start=1):
                root = ElementTree.fromstring(archive.read(name))
                text = normalize_extracted_text("\n".join(
                    node.text or "" for node in root.findall(".//a:t", self.namespaces)
                ))
                if text:
                    pages.append(ExtractedPage(number, text))
        return ParsedResourceDocument(
            pages=pages,
            sections=sections_from_pages(pages),
            tables=[],
            visuals=[],
            parser_name=self.name,
            ocr_required=not pages,
        )


class XLSXResourceParser(ResourceDocumentParser):
    name = "openpyxl"
    extensions = {".xlsx"}

    def parse(self, path: Path) -> ParsedResourceDocument:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        pages: list[ExtractedPage] = []
        sections: list[ExtractedSection] = []
        tables: list[ExtractedTable] = []
        for number, sheet in enumerate(workbook.worksheets, start=1):
            rows = [
                [normalize_extracted_text(value) for value in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            rows = [row for row in rows if any(row)]
            if not rows:
                continue
            raw = "\n".join(" | ".join(row) for row in rows)
            pages.append(ExtractedPage(number, f"{sheet.title}\n\n{raw}"))
            sections.append(ExtractedSection(
                title=sheet.title,
                content=raw,
                page_start=number,
                page_end=number,
                sort_order=len(sections),
                heading_path=[sheet.title],
            ))
            tables.append(ExtractedTable(number, sheet.title, rows[0], rows[1:], raw))
        workbook.close()
        return ParsedResourceDocument(
            pages=pages,
            sections=sections,
            tables=tables,
            visuals=[],
            parser_name=self.name,
            ocr_required=False,
        )


class TextResourceParser(ResourceDocumentParser):
    name = "text"
    extensions = {".txt", ".md", ".csv", ".log"}

    def parse(self, path: Path) -> ParsedResourceDocument:
        text = normalize_extracted_text(path.read_text(encoding="utf-8", errors="replace"))
        pages = [ExtractedPage(1, text)] if text else []
        return ParsedResourceDocument(
            pages=pages,
            sections=sections_from_pages(pages),
            tables=[],
            visuals=[],
            parser_name=self.name,
            ocr_required=False,
        )


PARSERS = [
    PDFResourceParser(),
    DOCXResourceParser(),
    PPTXResourceParser(),
    XLSXResourceParser(),
    TextResourceParser(),
]


def parser_for(path: Path) -> ResourceDocumentParser:
    parser = next((item for item in PARSERS if item.supports(path)), None)
    if not parser:
        raise ValueError(f"Unsupported knowledge document type: {path.suffix or 'unknown'}")
    return parser


def parse_resource_document(path: Path) -> ParsedResourceDocument:
    return parser_for(path).parse(path)


def extract_pages(path: Path) -> list[ExtractedPage]:
    return parse_resource_document(path).pages


def _split_content(content: str, max_characters: int, min_characters: int) -> list[str]:
    remaining = normalize_extracted_text(content)
    parts: list[str] = []
    while len(remaining) > max_characters:
        split_at = remaining.rfind(". ", 0, max_characters)
        split_at = split_at + 1 if split_at > min_characters else max_characters
        parts.append(remaining[:split_at].strip())
        remaining = remaining[split_at:].strip()
    if remaining:
        parts.append(remaining)
    return parts


def build_chunks(
    pages: list[ExtractedPage],
    *,
    sections: list[ExtractedSection] | None = None,
    tables: list[ExtractedTable] | None = None,
    maximum_tokens: int = DEFAULT_MAX_CHUNK_TOKENS,
    minimum_tokens: int = DEFAULT_MIN_CHUNK_TOKENS,
) -> list[ExtractedChunk]:
    maximum_characters = max(600, int(maximum_tokens) * 4)
    minimum_characters = max(100, int(minimum_tokens) * 4)
    source_sections = sections if sections is not None else sections_from_pages(pages)
    raw_chunks: list[tuple[int, int, str, list[str], str, str]] = []
    for section in source_sections:
        for part in _split_content(section.content, maximum_characters, minimum_characters):
            raw_chunks.append((
                section.page_start,
                section.page_end,
                section.title,
                section.heading_path,
                "Text",
                part,
            ))
    for table in tables or []:
        for part in _split_content(table.raw_representation, maximum_characters, minimum_characters):
            raw_chunks.append((
                table.page,
                table.page,
                table.title,
                [table.title],
                "Table",
                part,
            ))

    chunks: list[ExtractedChunk] = []
    for index, (start, end, heading, heading_path, chunk_type, content) in enumerate(raw_chunks):
        digest = hashlib.sha256(content.encode("utf-8")).hexdigest()
        page_label = str(start) if start == end else f"{start}-{end}"
        chunks.append(ExtractedChunk(
            index=index,
            page_start=start,
            page_end=end,
            heading=heading[:1000],
            content=content,
            content_hash=digest,
            token_count=max(1, round(len(content) / 4)),
            heading_path=heading_path,
            chunk_type=chunk_type,
            source_reference=f"{heading or 'Document'} · page {page_label}",
        ))
    return chunks
